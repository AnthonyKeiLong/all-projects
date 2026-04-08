"""
Extract and normalise plain text from .docx and .doc files.

Only document body text is compared (no metadata, styles, headers/footers,
comments, or revision marks).

  .docx  – read with python-docx (no external dependencies).
  .doc   – read via Word COM automation (requires pywin32 + Microsoft Word
           installed on Windows).
"""
import contextlib
import re
from pathlib import Path
from typing import Generator, Optional

try:
    from docx import Document
    from docx.opc.exceptions import PackageNotFoundError
except ImportError:  # pragma: no cover
    raise ImportError("python-docx is required. Install with: pip install python-docx")

# Optional: pywin32 for legacy .doc support (Windows only).
_WIN32COM_AVAILABLE = False
try:
    import pythoncom          # noqa: F401
    import win32com.client    # noqa: F401
    _WIN32COM_AVAILABLE = True
except ImportError:
    pass


# ── Shared normalisation ───────────────────────────────────────────────────────

def _normalise(text: str) -> str:
    """
    Normalise extracted raw text:
    - Unify line endings to ``\\n``.
    - Collapse horizontal whitespace (spaces/tabs) to a single space.
    - Collapse 3+ consecutive blank lines to two blank lines.
    - Strip leading/trailing whitespace.
    """
    text = re.sub(r"\r\n|\r", "\n", text)
    text = re.sub(r"[ \t]+", " ", text)
    text = re.sub(r"\n{3,}", "\n\n", text)
    return text.strip()


# ── .docx extraction ───────────────────────────────────────────────────────────

def _extract_text_docx(filepath: str) -> str:
    """Extract and normalise text from a .docx file using python-docx."""
    try:
        doc = Document(str(filepath))
    except PackageNotFoundError:
        raise ValueError(
            f"Cannot read '{filepath}': not a valid .docx file or file is corrupted."
        )
    except Exception as exc:
        raise ValueError(f"Cannot read '{filepath}': {exc}") from exc

    lines: list[str] = []

    # Main body paragraphs
    for para in doc.paragraphs:
        lines.append(para.text.strip())

    # Table cells (row-by-row, left-to-right; skip merged-cell duplicates)
    for table in doc.tables:
        for row in table.rows:
            deduped: list[str] = []
            for cell in row.cells:
                text = cell.text.strip()
                if not deduped or deduped[-1] != text:
                    deduped.append(text)
            for cell_text in deduped:
                if cell_text:
                    lines.append(cell_text)

    return _normalise("\n".join(lines))


# ── .doc extraction ────────────────────────────────────────────────────────────

def _extract_text_doc(filepath: str, word_app=None) -> str:
    """
    Extract text from a legacy .doc file via Word COM automation.

    Requires ``pywin32`` and Microsoft Word installed on Windows.

    Args:
        word_app: An existing ``win32com.client.Dispatch("Word.Application")``
                  instance to reuse.  When ``None`` a temporary instance is
                  created and quit after this single call (slower; prefer
                  passing a shared instance for batch processing).
    Raises:
        ValueError: if pywin32 is missing, Word is not installed, or the file
                    cannot be opened.
    """
    if not _WIN32COM_AVAILABLE:
        raise ValueError(
            f"Reading .doc files requires pywin32 and Microsoft Word. "
            "Install pywin32 with: pip install pywin32"
        )

    import pythoncom
    import win32com.client as _wc

    manage_word = word_app is None
    if manage_word:
        pythoncom.CoInitialize()
        word_app = _wc.Dispatch("Word.Application")
        word_app.Visible = False
        word_app.DisplayAlerts = 0   # wdAlertsNone

    try:
        abs_path = str(Path(filepath).resolve())
        doc = word_app.Documents.Open(
            abs_path,
            ReadOnly=True,
            AddToRecentFiles=False,
        )
        try:
            raw = doc.Content.Text
        finally:
            doc.Close(SaveChanges=False)
    except Exception as exc:
        raise ValueError(f"Cannot read '{filepath}': {exc}") from exc
    finally:
        if manage_word:
            try:
                word_app.Quit()
            except Exception:
                pass
            finally:
                pythoncom.CoUninitialize()

    return _normalise(raw)


# ── Word session context manager ───────────────────────────────────────────────

@contextlib.contextmanager
def word_session() -> Generator[Optional[object], None, None]:
    """
    Context manager that yields a reusable ``Word.Application`` COM instance
    for batch ``.doc`` processing.

    Yields ``None`` when pywin32 is unavailable (so callers degrade gracefully).

    Usage::

        with word_session() as app:
            for path in doc_files:
                text = extract_text(path, word_app=app)
    """
    if not _WIN32COM_AVAILABLE:
        yield None
        return

    import pythoncom
    import win32com.client as _wc

    app = None
    try:
        pythoncom.CoInitialize()
        try:
            app = _wc.Dispatch("Word.Application")
            app.Visible = False
            app.DisplayAlerts = 0
        except Exception:
            # Word is not installed or COM failed; individual .doc files will
            # each receive a STATUS_ERROR result — .docx scanning is unaffected.
            app = None
        yield app
    finally:
        if app is not None:
            try:
                app.Quit()
            except Exception:
                pass
        try:
            pythoncom.CoUninitialize()
        except Exception:
            pass


# ── Public API ─────────────────────────────────────────────────────────────────

def extract_text(filepath: str, word_app=None) -> str:
    """
    Return normalised plain text from a ``.docx`` or ``.doc`` file.

    Normalisation: trim, collapse whitespace, unify line endings.
    Ignored: metadata, styles, headers/footers, revision marks.

    Args:
        filepath:  Path to the Word file (absolute or relative).
        word_app:  Optional reusable ``Word.Application`` COM instance for
                   batch ``.doc`` extraction.  Ignored for ``.docx`` files.
    Raises:
        ValueError: if the file cannot be read or the extension is unsupported.
    """
    ext = Path(filepath).suffix.lower()
    if ext == ".docx":
        return _extract_text_docx(filepath)
    if ext == ".doc":
        return _extract_text_doc(filepath, word_app=word_app)
    raise ValueError(
        f"Unsupported file format '{ext}'. Only .docx and .doc are supported."
    )
