"""
Unit tests for the core comparison logic.

These tests require no GUI and only a standard Python environment + python-docx.
Run with:  pytest tests/ -v
"""
import os
import tempfile

import pytest
from docx import Document


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------

def _make_docx(path: str, paragraphs: list[str]) -> None:
    """Write a minimal .docx with the given paragraphs to *path*."""
    doc = Document()
    for text in paragraphs:
        doc.add_paragraph(text)
    doc.save(path)


# ---------------------------------------------------------------------------
# core.extractor
# ---------------------------------------------------------------------------

class TestExtractText:

    def test_basic_paragraphs(self, tmp_path):
        from core.extractor import extract_text
        fp = str(tmp_path / "a.docx")
        _make_docx(fp, ["Hello world", "Second paragraph"])
        result = extract_text(fp)
        assert "Hello world" in result
        assert "Second paragraph" in result

    def test_whitespace_normalisation(self, tmp_path):
        from core.extractor import extract_text
        fp = str(tmp_path / "ws.docx")
        _make_docx(fp, ["  lots   of   spaces  "])
        result = extract_text(fp)
        # Horizontal whitespace should be collapsed
        assert "lots of spaces" in result

    def test_empty_document(self, tmp_path):
        from core.extractor import extract_text
        fp = str(tmp_path / "empty.docx")
        _make_docx(fp, [])
        result = extract_text(fp)
        assert result == ""

    def test_invalid_file_raises(self, tmp_path):
        from core.extractor import extract_text
        fp = str(tmp_path / "bad.docx")
        fp_obj = tmp_path / "bad.docx"
        fp_obj.write_bytes(b"not a real docx file")
        with pytest.raises(ValueError):
            extract_text(fp)

    def test_table_text_extracted(self, tmp_path):
        from core.extractor import extract_text
        fp = str(tmp_path / "table.docx")
        doc = Document()
        tbl = doc.add_table(rows=2, cols=2)
        tbl.cell(0, 0).text = "Alpha"
        tbl.cell(0, 1).text = "Beta"
        tbl.cell(1, 0).text = "Gamma"
        tbl.cell(1, 1).text = "Delta"
        doc.save(fp)
        result = extract_text(fp)
        for word in ("Alpha", "Beta", "Gamma", "Delta"):
            assert word in result


# ---------------------------------------------------------------------------
# core.comparator
# ---------------------------------------------------------------------------

class TestComparatorHash:

    def test_identical_texts(self):
        from core.comparator import compare_texts
        assert compare_texts("hello", "hello") == "identical"

    def test_different_texts(self):
        from core.comparator import compare_texts
        assert compare_texts("hello", "world") == "different"

    def test_empty_vs_empty(self):
        from core.comparator import compare_texts
        assert compare_texts("", "") == "identical"

    def test_whitespace_sensitivity(self):
        from core.comparator import compare_texts
        # After normalisation by extractor, "a b" vs "a  b" would be the same;
        # but comparator itself is hash-based and byte-exact.
        assert compare_texts("a b", "a  b") == "different"


class TestSideBySideDiff:

    def test_equal_texts(self):
        from core.comparator import compute_side_by_side_diff
        la, lb = compute_side_by_side_diff("line1\nline2", "line1\nline2")
        assert all(t == "equal" for _, t in la)
        assert all(t == "equal" for _, t in lb)

    def test_added_line(self):
        from core.comparator import compute_side_by_side_diff
        la, lb = compute_side_by_side_diff("line1", "line1\nline2")
        assert len(la) == len(lb)
        # The second entry in lb should be "insert"
        assert any(t == "insert" for _, t in lb)

    def test_removed_line(self):
        from core.comparator import compute_side_by_side_diff
        la, lb = compute_side_by_side_diff("line1\nline2", "line1")
        assert len(la) == len(lb)
        assert any(t == "delete" for _, t in la)

    def test_replaced_line(self):
        from core.comparator import compute_side_by_side_diff
        la, lb = compute_side_by_side_diff("foo", "bar")
        assert len(la) == len(lb)
        assert any(t == "replace" for _, t in la)
        assert any(t == "replace" for _, t in lb)

    def test_output_lengths_equal(self):
        from core.comparator import compute_side_by_side_diff
        text_a = "a\nb\nc\nd"
        text_b = "a\nX\nY\nd\ne"
        la, lb = compute_side_by_side_diff(text_a, text_b)
        assert len(la) == len(lb)


# ---------------------------------------------------------------------------
# core.scanner
# ---------------------------------------------------------------------------

class TestScanner:

    def test_scan_empty_folder(self, tmp_path):
        from core.scanner import scan_folder
        result = scan_folder(str(tmp_path))
        assert result == {}

    def test_scan_finds_docx(self, tmp_path):
        from core.scanner import scan_folder
        fp = tmp_path / "report.docx"
        _make_docx(str(fp), ["test"])
        result = scan_folder(str(tmp_path))
        assert "report.docx" in result

    def test_scan_skips_temp_files(self, tmp_path):
        from core.scanner import scan_folder
        fp = tmp_path / "~$report.docx"
        fp.write_bytes(b"lock")
        result = scan_folder(str(tmp_path))
        assert "~$report.docx" not in result

    def test_scan_finds_doc(self, tmp_path):
        """Legacy .doc files should be discovered regardless of content."""
        from core.scanner import scan_folder
        fp = tmp_path / "legacy.doc"
        fp.write_bytes(b"dummy")   # scanner checks extension only
        result = scan_folder(str(tmp_path))
        assert "legacy.doc" in result

    def test_scan_skips_doc_temp_files(self, tmp_path):
        from core.scanner import scan_folder
        fp = tmp_path / "~$legacy.doc"
        fp.write_bytes(b"lock")
        result = scan_folder(str(tmp_path))
        assert "~$legacy.doc" not in result

    def test_scan_ignores_non_word_extensions(self, tmp_path):
        from core.scanner import scan_folder
        (tmp_path / "notes.txt").write_text("text")
        (tmp_path / "report.pdf").write_bytes(b"pdf")
        result = scan_folder(str(tmp_path))
        assert result == {}

    def test_scan_recursive(self, tmp_path):
        from core.scanner import scan_folder
        sub = tmp_path / "sub"
        sub.mkdir()
        _make_docx(str(sub / "nested.docx"), ["nested"])
        result = scan_folder(str(tmp_path), recursive=True)
        assert "sub/nested.docx" in result

    def test_scan_non_recursive(self, tmp_path):
        from core.scanner import scan_folder
        sub = tmp_path / "sub"
        sub.mkdir()
        _make_docx(str(sub / "nested.docx"), ["nested"])
        result = scan_folder(str(tmp_path), recursive=False)
        assert "sub/nested.docx" not in result

    def test_match_files_by_filename(self, tmp_path):
        from core.scanner import match_files
        folder_a = tmp_path / "a"
        folder_b = tmp_path / "b"
        folder_a.mkdir(); folder_b.mkdir()
        _make_docx(str(folder_a / "doc.docx"), ["version A"])
        _make_docx(str(folder_b / "doc.docx"), ["version B"])
        pairs = match_files(str(folder_a), str(folder_b), match_by="filename")
        assert len(pairs) == 1
        key, pa, pb = pairs[0]
        assert key == "doc.docx"
        assert pa is not None and pb is not None

    def test_match_files_only_one_side(self, tmp_path):
        from core.scanner import match_files
        folder_a = tmp_path / "a"
        folder_b = tmp_path / "b"
        folder_a.mkdir(); folder_b.mkdir()
        _make_docx(str(folder_a / "only_a.docx"), ["only in A"])
        pairs = match_files(str(folder_a), str(folder_b), match_by="filename")
        assert len(pairs) == 1
        _, pa, pb = pairs[0]
        assert pa is not None
        assert pb is None


# ---------------------------------------------------------------------------
# core.extractor – dispatch tests
# ---------------------------------------------------------------------------

class TestExtractDispatch:

    def test_unsupported_extension_raises(self, tmp_path):
        from core.extractor import extract_text
        fp = str(tmp_path / "file.rtf")
        (tmp_path / "file.rtf").write_text("rtf content")
        with pytest.raises(ValueError, match="Unsupported"):
            extract_text(fp)

    def test_doc_without_pywin32_raises_value_error(self, tmp_path, monkeypatch):
        """Simulate pywin32 absent: should raise ValueError with helpful message."""
        import core.extractor as ext_mod
        monkeypatch.setattr(ext_mod, "_WIN32COM_AVAILABLE", False)
        fp = str(tmp_path / "legacy.doc")
        (tmp_path / "legacy.doc").write_bytes(b"dummy")
        with pytest.raises(ValueError, match="pywin32"):
            ext_mod.extract_text(fp)


# ---------------------------------------------------------------------------
# core.exporter
# ---------------------------------------------------------------------------

class TestExporter:

    def test_export_csv_creates_file(self, tmp_path):
        from core.exporter import export_csv
        from core.models   import ResultItem, STATUS_DIFFERENT, DECISION_PENDING
        items = [
            ResultItem(key="file.docx", path_a="/a/file.docx",
                       path_b="/b/file.docx", status=STATUS_DIFFERENT,
                       decision=DECISION_PENDING),
        ]
        out = str(tmp_path / "report.csv")
        n = export_csv(out, items)
        assert n == 1
        assert os.path.isfile(out)
        content = open(out, encoding="utf-8-sig").read()
        assert "file.docx" in content
        assert "Different" in content


# ---------------------------------------------------------------------------
# core.session
# ---------------------------------------------------------------------------

class TestSession:

    def test_round_trip(self, tmp_path):
        from core.session  import save_session, load_session
        from core.models   import ResultItem, STATUS_DIFFERENT, DECISION_REVIEWED
        items = [
            ResultItem(key="x.docx", path_a="/a/x.docx", path_b="/b/x.docx",
                       status=STATUS_DIFFERENT, decision=DECISION_REVIEWED),
        ]
        fp = str(tmp_path / "session.json")
        save_session(fp, "/folder/a", "/folder/b", True, "filename", items)

        data = load_session(fp)
        assert data["folder_a"] == "/folder/a"
        assert data["folder_b"] == "/folder/b"
        assert data["recursive"] is True
        assert data["match_by"]  == "filename"
        assert len(data["results"]) == 1
        restored = data["results"][0]
        assert restored.key      == "x.docx"
        assert restored.decision == DECISION_REVIEWED
